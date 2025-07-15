# src/app/tabs/export_tab.py

import streamlit as st
from io import BytesIO
from pathlib import Path
from datetime import datetime
import re

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# -----------------------------
# Utility Functions
# -----------------------------

def slugify(text):
    return re.sub(r'[\W_]+', '_', text.strip().lower())

def get_filename(title, ext):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{timestamp}_{slugify(title)}.{ext}"

def write_markdown(title, metadata, content):
    return (
        f"# {title}\n\n"
        f"**Tone**: {metadata.get('tone', '-')}\n\n"
        f"**Audience**: {metadata.get('audience', '-')}\n\n"
        f"**Call to Action**: {metadata.get('cta', '-')}\n\n"
        "---\n\n"
        f"{content}"
    )

def generate_txt(markdown_text):
    return markdown_text.encode("utf-8")

def generate_md(markdown_text):
    return markdown_text.encode("utf-8")

def generate_docx(title, metadata, content):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Tone: {metadata.get('tone', '-')}")
    doc.add_paragraph(f"Audience: {metadata.get('audience', '-')}")
    doc.add_paragraph(f"Call to Action: {metadata.get('cta', '-')}")
    doc.add_paragraph("")  # spacer
    doc.add_paragraph(content)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def generate_pdf(title, metadata, content):
    """
    Generate a PDF using ReportLab Platypus.
    Returns PDF bytes suitable for st.download_button.
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=40, leftMargin=40,
                            topMargin=60, bottomMargin=40)

    styles = getSampleStyleSheet()
    
    # Title style
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Title'],
        fontSize=20,
        textColor=colors.HexColor("#333333"),
        spaceAfter=12
    )
    
    normal_style = styles["Normal"]
    normal_style.fontSize = 12
    normal_style.leading = 16

    elements = []

    # Title
    elements.append(Paragraph(title, title_style))
    elements.append(Spacer(1, 12))

    # Metadata Table
    meta_data = [
        ["Tone:", metadata.get("tone", "-")],
        ["Audience:", metadata.get("audience", "-")],
        ["Call to Action:", metadata.get("cta", "-")],
    ]

    table = Table(meta_data, colWidths=[100, 350])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.darkblue),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 11),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("BACKGROUND", (0, 1), (-1, -1), colors.whitesmoke),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 12))

    # Content
    safe_content = content.replace('\n', '<br/>')
    elements.append(Paragraph(safe_content, normal_style))

    doc.build(elements)
    buffer.seek(0)
    return buffer.read()

# -----------------------------
# Main UI Function
# -----------------------------

def render_export_tab():
    st.header("Step 8: Export Final Content")
    st.info("Download your polished content in Markdown, PDF, Word, or Plain Text formats.")

    polished_drafts = st.session_state.get("final_export_ready", [])

    if not polished_drafts:
        st.warning("⚠️ No drafts selected for export. Please complete Step 7: Polish Content.")
        return

    for i, draft in enumerate(polished_drafts):
        st.markdown("---")
        st.subheader(f"📄 {draft.get('title', f'Draft {i+1}')}")

        title = draft.get("title", f"Draft_{i+1}")
        content = draft.get("polished", "")
        brief = draft.get("brief", {})
        metadata = {
            "tone": draft.get("tone", "Professional"),
            "audience": draft.get("audience", "Business Leaders"),
            "cta": brief.get("cta", "")
        }

        md_content = write_markdown(title, metadata, content)

        # Preview Section
        with st.expander("🔍 Preview Polished Content", expanded=False):
            st.markdown(f"**Tone:** {metadata['tone']}")
            st.markdown(f"**Audience:** {metadata['audience']}")
            st.markdown(f"**Call to Action:** {metadata['cta']}")
            st.markdown("---")
            st.markdown(content)

        # --------------------------
        # Markdown Download
        # --------------------------
        if st.button(f"Generate Markdown (.md) for {title}", key=f"btn_md_{i}"):
            md_bytes = generate_md(md_content)
            st.download_button(
                label="⬇️ Download Markdown (.md)",
                data=md_bytes,
                file_name=get_filename(title, "md"),
                mime="text/markdown",
                key=f"dl_md_{i}"
            )

        # --------------------------
        # Text Download
        # --------------------------
        if st.button(f"Generate Plain Text (.txt) for {title}", key=f"btn_txt_{i}"):
            txt_bytes = generate_txt(md_content)
            st.download_button(
                label="⬇️ Download Plain Text (.txt)",
                data=txt_bytes,
                file_name=get_filename(title, "txt"),
                mime="text/plain",
                key=f"dl_txt_{i}"
            )

        # --------------------------
        # Word DOCX Download
        # --------------------------
        if st.button(f"Generate Word (.docx) for {title}", key=f"btn_docx_{i}"):
            docx_bytes = generate_docx(title, metadata, content)
            st.download_button(
                label="⬇️ Download Word (.docx)",
                data=docx_bytes,
                file_name=get_filename(title, "docx"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_docx_{i}"
            )

        # --------------------------
        # PDF Download
        # --------------------------
        if st.button(f"Generate PDF (.pdf) for {title}", key=f"btn_pdf_{i}"):
            pdf_bytes = generate_pdf(title, metadata, content)
            st.download_button(
                label="⬇️ Download PDF (.pdf)",
                data=pdf_bytes,
                file_name=get_filename(title, "pdf"),
                mime="application/pdf",
                key=f"dl_pdf_{i}"
            )